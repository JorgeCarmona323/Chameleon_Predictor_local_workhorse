"""
Convert writeup_2026-03-18.md to a formatted Word document.
Usage: python scripts/export_docx.py
Output: docs/Carmona_Chameleon_Predictor_Report_2026-03-18.docx
"""

import re
from pathlib import Path
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

MD_PATH  = Path(__file__).parent.parent / "docs" / "writeup_2026-03-18.md"
OUT_PATH = Path(__file__).parent.parent / "docs" / "Carmona_Chameleon_Predictor_Report_2026-03-18.docx"
FIG_DIR  = Path(__file__).parent.parent / "results" / "figures"

# ── Helpers ────────────────────────────────────────────────────────────────────

def set_font(run, name="Calibri", size=11, bold=False, italic=False, color=None):
    run.font.name = name
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    if color:
        run.font.color.rgb = RGBColor(*color)

def add_heading(doc, text, level):
    p = doc.add_heading(text, level=level)
    # Ensure heading font is clean
    for run in p.runs:
        run.font.name = "Calibri"
    return p

# Unicode superscript digit map → normal digit string
_SUP_CHARS = str.maketrans("⁰¹²³⁴⁵⁶⁷⁸⁹", "0123456789")
_SUP_RE = re.compile(r'[⁰¹²³⁴⁵⁶⁷⁸⁹]+')

def _add_runs(p, text, bold=False, italic=False, name="Calibri", size=11, color=None, code=False):
    """Split text on superscript unicode runs and add each segment as a proper run."""
    pos = 0
    for m in _SUP_RE.finditer(text):
        # Normal text before the superscript
        if m.start() > pos:
            run = p.add_run(text[pos:m.start()])
            set_font(run, name=name, size=size, bold=bold, italic=italic, color=color)
            if code:
                run.font.name = "Courier New"
        # Superscript run
        sup_run = p.add_run(m.group().translate(_SUP_CHARS))
        set_font(sup_run, name="Calibri", size=8, bold=False, italic=False)
        sup_run.font.superscript = True
        pos = m.end()
    # Remaining text after last superscript
    if pos < len(text):
        run = p.add_run(text[pos:])
        set_font(run, name=name, size=size, bold=bold, italic=italic, color=color)
        if code:
            run.font.name = "Courier New"

def add_paragraph(doc, text, style="Normal"):
    p = doc.add_paragraph(style=style)
    # Split on bold (**text**) and inline code (`text`) first
    parts = re.split(r'(\*\*[^*]+\*\*|`[^`]+`)', text)
    for part in parts:
        if part.startswith("**") and part.endswith("**"):
            _add_runs(p, part[2:-2], bold=True)
        elif part.startswith("`") and part.endswith("`"):
            _add_runs(p, part[1:-1], name="Courier New", size=9,
                      color=(0x2E, 0x86, 0xC1), code=True)
        else:
            _add_runs(p, part)
    return p

def add_table(doc, header_row, data_rows):
    n_cols = len(header_row)
    table = doc.add_table(rows=1 + len(data_rows), cols=n_cols)
    table.style = "Table Grid"

    # Header
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(header_row):
        hdr_cells[i].text = h.strip()
        for para in hdr_cells[i].paragraphs:
            for run in para.runs:
                set_font(run, bold=True, size=10)
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        # Header bg: light blue-grey
        tc = hdr_cells[i]._tc
        tcPr = tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"), "D6E4F0")
        tcPr.append(shd)

    # Data rows
    for r_idx, row in enumerate(data_rows):
        cells = table.rows[r_idx + 1].cells
        for c_idx, cell_text in enumerate(row):
            cells[c_idx].text = cell_text.strip()
            for para in cells[c_idx].paragraphs:
                for run in para.runs:
                    set_font(run, size=10)
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()  # spacing after table
    return table

def add_code_block(doc, code_text):
    p = doc.add_paragraph(style="Normal")
    run = p.add_run(code_text)
    set_font(run, name="Courier New", size=9)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    for side in ("top", "left", "bottom", "right"):
        bdr = OxmlElement(f"w:{side}")
        bdr.set(qn("w:val"), "single")
        bdr.set(qn("w:sz"), "4")
        bdr.set(qn("w:space"), "4")
        bdr.set(qn("w:color"), "AAAAAA")
        pBdr.append(bdr)
    pPr.append(pBdr)
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), "F2F3F4")
    pPr.append(shd)
    return p

def try_insert_figure(doc, fig_name, caption):
    """Insert figure image if it exists, else insert placeholder."""
    # Map markdown image links to actual filenames
    candidates = [
        FIG_DIR / fig_name,
        FIG_DIR / (fig_name.replace("figures/", "")),
    ]
    for path in candidates:
        if path.exists():
            try:
                doc.add_picture(str(path), width=Inches(5.5))
                last_para = doc.paragraphs[-1]
                last_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                cap = doc.add_paragraph(caption)
                cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in cap.runs:
                    set_font(run, italic=True, size=9)
                return
            except Exception:
                pass
    # Placeholder
    p = doc.add_paragraph(f"[Figure: {caption}]")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in p.runs:
        set_font(run, italic=True, color=(0x99, 0x99, 0x99))

# ── Parser ─────────────────────────────────────────────────────────────────────

def parse_md_table(lines):
    """Parse markdown table lines into (header, rows)."""
    rows = []
    for line in lines:
        cells = [c for c in line.strip().split("|") if c.strip() != ""]
        rows.append(cells)
    if not rows:
        return [], []
    # Second row is separator (---), skip it
    header = rows[0] if rows else []
    data   = [r for r in rows[2:] if r]
    return header, data

def build_docx(md_path, out_path):
    doc = Document()

    # Page margins
    for section in doc.sections:
        section.top_margin    = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin   = Cm(3.0)
        section.right_margin  = Cm(3.0)

    # Default paragraph style
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    text = md_path.read_text(encoding="utf-8")
    lines = text.splitlines()

    i = 0
    while i < len(lines):
        line = lines[i]

        # ── Skip horizontal rules ────────────────────────────────────────────
        if re.match(r"^---+\s*$", line):
            i += 1
            continue

        # ── Headings ─────────────────────────────────────────────────────────
        m = re.match(r"^(#{1,4})\s+(.*)", line)
        if m:
            level = len(m.group(1))
            text_ = m.group(2).strip()
            # Title (level 1) — treat as document title
            if level == 1:
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = p.add_run(text_)
                set_font(run, size=18, bold=True)
            else:
                add_heading(doc, text_, level=level - 1)
            i += 1
            continue

        # ── Author / date line (bold+italic line after title) ─────────────
        if line.startswith("**") and line.endswith("**") and "March" in line:
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(line.strip("*"))
            set_font(run, bold=True, italic=True, size=11)
            i += 1
            continue

        # ── Markdown tables ──────────────────────────────────────────────────
        if line.startswith("|"):
            table_lines = []
            while i < len(lines) and lines[i].startswith("|"):
                table_lines.append(lines[i])
                i += 1
            header, data = parse_md_table(table_lines)
            if header:
                add_table(doc, header, data)
            continue

        # ── Code blocks ──────────────────────────────────────────────────────
        if line.startswith("```"):
            i += 1
            code_lines = []
            while i < len(lines) and not lines[i].startswith("```"):
                code_lines.append(lines[i])
                i += 1
            i += 1  # closing ```
            add_code_block(doc, "\n".join(code_lines))
            continue

        # ── Images ───────────────────────────────────────────────────────────
        m = re.match(r"!\[([^\]]*)\]\(([^)]+)\)", line)
        if m:
            caption = m.group(1)
            src     = m.group(2)
            # Extract just filename
            fig_name = Path(src).name
            try_insert_figure(doc, fig_name, caption)
            i += 1
            continue

        # ── Bullet lists ──────────────────────────────────────────────────────
        m = re.match(r"^[-*]\s+(.*)", line)
        if m:
            p = add_paragraph(doc, m.group(1), style="List Bullet")
            i += 1
            continue

        # ── Blank line ────────────────────────────────────────────────────────
        if not line.strip():
            i += 1
            continue

        # ── Normal paragraph ─────────────────────────────────────────────────
        add_paragraph(doc, line)
        i += 1

    doc.save(out_path)
    print(f"Saved: {out_path}")

if __name__ == "__main__":
    build_docx(MD_PATH, OUT_PATH)
